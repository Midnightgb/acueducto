# FUNCION GENERAR TOKEN
from datetime import datetime, timedelta
from jose import jwt
from docx import Document
from fpdf import FPDF
from models import Token, Usuario, Empresa, Vivienda

SECRET_KEY = "sd45g4f45SWFGVHHuoyiad4F5SFD65V4SFDVOJWNHACUfwghdfvcguDCwfghezxhAzAKHGFBJYTFdkjfghtjkdgb"


def generar_token(usuario_id: str):
    expiration = datetime.utcnow() + timedelta(hours=1)  # Token expira en 1 hora
    payload = {"sub": usuario_id, "exp": expiration}
    token = jwt.encode(payload, SECRET_KEY, algorithm="HS256")
    return token

# FUNCION VERIFICAR TOKEN


def verificar_token(token: str, db):
    if token is None:
        return None
    consultar = db.query(Token).filter(Token.token == token).first()
    if consultar:
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms="HS256")
            return payload.get("sub")
        except jwt.ExpiredSignatureError:  # Token ha expirado
            deleteToken = db.query(Token).filter(Token.token == token).first()
            if deleteToken:
                db.delete(deleteToken)
                db.commit()
            return None
        except jwt.JWTError:  # Token inválido
            return None
    else:  # Token no válido
        return None


# FUNCION PARA OBTENER EL ROL DE USUARIO
def get_rol(id_usuario, db):
    if id_usuario:
        usuario = db.query(Usuario).filter(
            Usuario.id_usuario == id_usuario).first()
        if usuario:
            return usuario.rol
        else:
            return None
    else:
        return None


def get_datos_vivienda(id_vivienda, db):
    if id_vivienda:
        vivienda = db.query(Vivienda).filter(
            Vivienda.id_inmueble == id_vivienda).first()
        if vivienda:
            datos_vivienda = {
                "id_vivienda": vivienda.id_inmueble,
                "id_usuario": vivienda.id_usuario,
                "direccion": vivienda.direccion,
                "estrato": vivienda.estrato,
                "uso": vivienda.uso,
                "numero_residentes": vivienda.numero_residentes,
            }
            return datos_vivienda
        else:
            return None
    else:
        return None


def get_viviendas(id_usuario, db):
    if id_usuario:
        viviendas = db.query(Vivienda).filter(
            Vivienda.id_usuario == id_usuario).all()
        if viviendas:
            return viviendas
        else:
            return None
    else:
        return None

# FUNCION PARA OBTENER LOS DATOS DE USUARIO


def get_datos_usuario(id_usuario, db):
    if id_usuario:
        usuario = db.query(Usuario).filter(
            Usuario.id_usuario == id_usuario).first()
        if usuario:
            datos_usuario = {
                "id_usuario": usuario.id_usuario,
                "nom_usuario": usuario.nom_usuario,
                "apellido_usuario": usuario.apellido_usuario,
                "num_doc": usuario.num_doc,
                "direccion": usuario.direccion,
                "municipio": usuario.municipio,
                "rol": usuario.rol,
                "estado": usuario.estado,
                "empresa": usuario.empresa,
                "correo": usuario.correo
                # Agrega otros campos del usuario
            }
            return datos_usuario
        else:
            return None
    else:
        return None

# CAMBIAR LOS CAMPOS "[]" POR VALORES DE FORMULARIO


def reemplazar_texto(docx_path, datos):
    document = Document(docx_path)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            texto = run.text
            for clave, valor in datos.items():
                texto = texto.replace(clave, valor)
            run.text = texto

    return document

# CONVIERTE EL DOCUMENTO DOCX A PDF


def convertir_a_pdf(docx_path, pdf_path):
    document = Document(docx_path)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    for paragraph in document.paragraphs:
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, paragraph.text)

    pdf.output(pdf_path)


# OBTENER DATOS EMPRESA:
def get_datos_empresa(id_empresa, db):
    if id_empresa:
        empresa = db.query(Empresa).filter(
            Empresa.id_empresa == id_empresa).first()
        if empresa:
            datos_empresa = {
                "id_empresa": empresa.id_empresa,
                "nom_empresa": empresa.nom_empresa,
                "tel_fijo": empresa.tel_fijo,
                "tel_cel": empresa.tel_cel,
                "email": empresa.email,
                "estado": empresa.estado,
            }
            return datos_empresa
        else:
            return None
    else:
        return None


# OBTERNER DATOS EMPRESA (TODAS LAS REGISTRADAS):

def get_datos_empresas(db) -> list[str]:
    nom_empresas = db.query(Empresa.nom_empresa).all()
    return [nombre[0] for nombre in nom_empresas]
